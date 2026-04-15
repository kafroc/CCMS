"""
ST document export routes.
- Template-based rendering with {placeholder} tokens
- Live preview endpoint (returns rendered markdown)
- Export to markdown / docx
"""
import json
import os
import re
from datetime import datetime, timezone
from typing import Optional

from fastapi import APIRouter, HTTPException, Depends, Query
from fastapi.responses import FileResponse, Response
from pydantic import BaseModel
from sqlalchemy.ext.asyncio import AsyncSession
from sqlmodel import select

from app.core.auth import get_current_user, get_current_admin
from app.core.config import settings
from app.core.database import get_db
from app.core.response import ok
from app.core.toe_permissions import get_accessible_toe
from app.models.user import User
from app.models.toe import TOE, TOEAsset
from app.models.threat import Threat, Assumption, OSP
from app.models.security import SecurityObjective, SFR, SFRLibrary, ObjectiveSFR, ThreatObjective, AssumptionObjective, OSPObjective
from app.models.test_case import TestCase
from app.models.system_setting import SystemSetting

router = APIRouter(prefix="/api", tags=["export"])

ST_TEMPLATE_FILENAME = "ST Template.md"
PROJECT_ROOT_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", ".."))

_SFR_DEPENDENCY_ID_RE = re.compile(r"\b[A-Z]{3}_[A-Z0-9]+(?:\.[A-Z0-9]+)+\b", re.IGNORECASE)
_SFR_DEPENDENCY_TOKEN_RE = re.compile(r"\[|\]|\(|\)|\bOR\b|[\n,;；，]+|\b[A-Z]{3}_[A-Z0-9]+(?:\.[A-Z0-9]+)+\b", re.IGNORECASE)


async def _get_toe(toe_id: str, user: User, db) -> TOE:
    return await get_accessible_toe(toe_id, user, db, writable=False)


def _get_st_template_file_path() -> str:
    return os.path.join(PROJECT_ROOT_DIR, ST_TEMPLATE_FILENAME)


def _write_st_template_file(content: str) -> None:
    file_path = _get_st_template_file_path()
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    with open(file_path, "w", encoding="utf-8") as file:
        file.write(content)


def _read_st_template_file() -> Optional[str]:
    file_path = _get_st_template_file_path()
    if not os.path.exists(file_path):
        return None
    with open(file_path, "r", encoding="utf-8") as file:
        return file.read()


# ── Default ST template ──────────────────────────────────────────

DEFAULT_ST_TEMPLATE = r"""# Security Target — {TOE Name}

**Product:** {TOE Name}  |  **Version:** {TOE Version}  |  **Type:** {TOE Type}

---

## 1. ST Introduction

{TOE Brief Intro}

## 2. TOE Description

### 2.1 Product Type
{TOE Type Desc}

### 2.2 Product Usage & Deployment
{TOE Usage}

### 2.3 Major Security Features
{Major Security Features}

### 2.4 Required Non-TOE Hardware / Software / Firmware
{Required Non-TOE}

## 3. TOE Scope & Boundaries

### 3.1 Physical Scope
{Physical Scope}

### 3.2 Logical Scope
{Logical Scope}

### 3.3 Hardware Interfaces
{HW Interfaces}

### 3.4 Software Interfaces
{SW Interfaces}

## 4. Protected Assets

{Assets Table}

## 5. Security Problem Definition

### 5.1 Threats

{Threats Table}

### 5.2 Assumptions

{Assumptions Table}

### 5.3 Organizational Security Policies

{OSP Table}

## 6. Security Objectives

### 6.1 Security Objectives for the TOE

{Objectives Table}

### 6.2 Security Objectives for the Operational Environment

{OE Table}

## 7. Security Functional Requirements (SFRs)

{SFR Table}

## 8. Security Objectives Rationale

### 8.1 TAO Objective Mapping

{TAO Objective Mapping}

### 8.2 Threat Objective Rationale

{Threat Objective rationale}

### 8.3 Assumption Objective Rationale

{Assumption Objective rationale}

### 8.4 OSP Objective Rationale

{OSP Objective rationale}

### 8.5 Threats → Objectives Mapping

{Threat Objective Mapping}

### 8.6 Assumptions → Objectives Mapping

{Assumption Objective Mapping}

### 8.7 OSP → Objectives Mapping

{OSP Objective Mapping}

### 8.8 Objectives → SFR Mapping

{Objective SFR Mapping}

### 8.9 Objective SFR Rationale

{Objective SFR rationale}

### 8.10 SFR Dependencies

{SFR Dependencies}

## 9. Test Cases

{Test Cases Table}
"""


# ── Rendering engine ─────────────────────────────────────────────

async def _load_all_data(toe_id: str, db: AsyncSession):
    """Load all TOE-related data for template rendering."""
    toe_res = await db.exec(select(TOE).where(TOE.id == toe_id))
    toe = toe_res.first()
    if not toe:
        raise HTTPException(404, "TOE not found")

    assets = (await db.exec(select(TOEAsset).where(TOEAsset.toe_id == toe_id, TOEAsset.deleted_at.is_(None)))).all()
    threats = (await db.exec(select(Threat).where(Threat.toe_id == toe_id, Threat.deleted_at.is_(None)).order_by(Threat.code))).all()
    assumptions = (await db.exec(select(Assumption).where(Assumption.toe_id == toe_id, Assumption.deleted_at.is_(None)).order_by(Assumption.code))).all()
    osps = (await db.exec(select(OSP).where(OSP.toe_id == toe_id, OSP.deleted_at.is_(None)).order_by(OSP.code))).all()
    objectives = (await db.exec(select(SecurityObjective).where(SecurityObjective.toe_id == toe_id, SecurityObjective.deleted_at.is_(None)).order_by(SecurityObjective.code))).all()
    sfrs = (await db.exec(select(SFR).where(SFR.toe_id == toe_id, SFR.deleted_at.is_(None)).order_by(SFR.sfr_id))).all()
    tests = (await db.exec(select(TestCase).where(TestCase.toe_id == toe_id, TestCase.deleted_at.is_(None)).order_by(TestCase.created_at))).all()

    threat_ids = [item.id for item in threats]
    assumption_ids = [item.id for item in assumptions]
    osp_ids = [item.id for item in osps]
    objective_ids = [item.id for item in objectives]
    sfr_library_ids = [item.sfr_library_id for item in sfrs if item.sfr_library_id]
    sfr_components = [item.sfr_id for item in sfrs if item.sfr_id]

    # Mapping tables
    threat_objectives = (await db.exec(select(ThreatObjective).where(ThreatObjective.threat_id.in_(threat_ids)))).all() if threat_ids else []
    assumption_objectives = (await db.exec(select(AssumptionObjective).where(AssumptionObjective.assumption_id.in_(assumption_ids)))).all() if assumption_ids else []
    osp_objectives = (await db.exec(select(OSPObjective).where(OSPObjective.osp_id.in_(osp_ids)))).all() if osp_ids else []
    objective_sfrs = (await db.exec(select(ObjectiveSFR).where(ObjectiveSFR.objective_id.in_(objective_ids)))).all() if objective_ids else []
    libraries_by_id = {
        item.id: item
        for item in ((await db.exec(select(SFRLibrary).where(SFRLibrary.id.in_(sfr_library_ids)))).all() if sfr_library_ids else [])
    }
    libraries_by_component = {
        item.sfr_component.upper(): item
        for item in ((await db.exec(select(SFRLibrary).where(SFRLibrary.sfr_component.in_(sfr_components)))).all() if sfr_components else [])
    }

    return {
        "toe": toe, "assets": assets, "threats": threats,
        "assumptions": assumptions, "osps": osps, "objectives": objectives,
        "sfrs": sfrs, "tests": tests,
        "threat_objectives": threat_objectives,
        "assumption_objectives": assumption_objectives,
        "osp_objectives": osp_objectives,
        "objective_sfrs": objective_sfrs,
        "libraries_by_id": libraries_by_id,
        "libraries_by_component": libraries_by_component,
    }


def _v(val: Optional[str]) -> str:
    """Return value or placeholder dash."""
    return val.strip() if val and val.strip() else "—"


def _md_table(headers: list[str], rows: list[list[str]]) -> str:
    """Build a markdown table."""
    if not rows:
        return "*（No records）*"
    header = "| " + " | ".join(headers) + " |"
    sep = "| " + " | ".join("---" for _ in headers) + " |"
    body = "\n".join("| " + " | ".join(re.sub(r'\r\n|\r|\n', ' ', c).replace("|", "\\|") for c in row) + " |" for row in rows)
    return header + "\n" + sep + "\n" + body


def _md_list(values: list[str]) -> str:
    """Build a markdown bullet list from a single column of values."""
    cleaned = [v.strip() for v in values if v and v.strip() and v.strip() != "—"]
    if not cleaned:
        return "—"
    if len(cleaned) == 1:
        return cleaned[0]
    return "\n".join(f"- {v}" for v in cleaned)


def _blank(val: Optional[str]) -> str:
    return val.strip() if val and val.strip() else ""


def _threat_definition(threat: Threat) -> str:
    parts = [part.strip() for part in [threat.threat_agent, threat.adverse_action, threat.assets_affected] if part and part.strip()]
    if not parts:
        return "—"
    return " / ".join(parts)


def _objective_groups(objectives: list[SecurityObjective]) -> tuple[list[SecurityObjective], list[SecurityObjective]]:
    obj_o = [o for o in objectives if (o.code or "").startswith("O.") or o.obj_type == "O"]
    obj_oe = [o for o in objectives if (o.code or "").startswith("OE.") or o.obj_type == "OE"]
    return obj_o, obj_oe


def _extract_dependency_ids(raw_dependency: Optional[str]) -> list[str]:
    if not raw_dependency:
        return []
    return re.findall(r'[A-Za-z]{2,4}_[A-Za-z]{2,4}(?:\.\d+)+', raw_dependency)


def _extract_dependency_groups(value: Optional[str]) -> list[list[str]]:
    text = (value or "").strip()
    if not text:
        return []
    tokens = [match.group(0) for match in _SFR_DEPENDENCY_TOKEN_RE.finditer(text.upper())]
    if not tokens:
        return []

    groups: list[list[str]] = []
    current_group: list[str] = []
    inside_brackets = False
    pending_or = False

    def flush_group() -> None:
        nonlocal current_group, pending_or
        if current_group:
            groups.append(current_group)
            current_group = []
        pending_or = False

    for token in tokens:
        if token in {"[", "("}:
            flush_group()
            inside_brackets = True
            continue
        if token in {"]", ")"}:
            flush_group()
            inside_brackets = False
            continue
        if token == "OR":
            pending_or = True
            continue
        if token in {",", ";", "；", "，", "\n", "\r\n"} or any(ch in token for ch in "\n,;；，"):
            if not inside_brackets:
                flush_group()
            continue
        if _SFR_DEPENDENCY_ID_RE.fullmatch(token):
            dependency_id = token.upper()
            if inside_brackets:
                if dependency_id not in current_group:
                    current_group.append(dependency_id)
                pending_or = False
                continue
            if not current_group:
                current_group = [dependency_id]
            elif pending_or:
                if dependency_id not in current_group:
                    current_group.append(dependency_id)
            else:
                flush_group()
                current_group = [dependency_id]
            pending_or = False

    flush_group()
    return groups


def _is_dependency_mapping(mapping_rationale: Optional[str]) -> bool:
    text = (mapping_rationale or "").strip()
    return "dependency supplement" in text.lower() or "auto-added as a dependency" in text.lower()


def _resolve_export_sfr_fields(
    sfr: SFR,
    libraries_by_id: dict[str, SFRLibrary],
    libraries_by_component: dict[str, SFRLibrary],
) -> dict[str, str]:
    library = libraries_by_id.get(sfr.sfr_library_id) if sfr.sfr_library_id else None
    if library is None and sfr.sfr_id:
        library = libraries_by_component.get(sfr.sfr_id.upper())
    return {
        "name": _v(sfr.sfr_name or (library.sfr_component_name if library else None)),
        "detail": _v(sfr.sfr_detail or (library.description if library else None)),
        "detail_raw": (sfr.sfr_detail or (library.description if library else None) or "").strip(),
        "dependency_raw": (sfr.dependency or (library.dependencies if library else None) or "").strip(),
    }


def _md_sfr_sections(sfrs: list[SFR], sfr_display_map: dict[str, dict[str, str]]) -> str:
    if not sfrs:
        return "*（No records）*"
    blocks: list[str] = []
    for sfr in sfrs:
        fields = sfr_display_map[sfr.id]
        title = f"{sfr.sfr_id} {fields['name']}".strip()
        detail = fields["detail_raw"] or fields["detail"]
        blocks.append(f"**{title}**\n{detail}")
    return "\n\n".join(blocks)


def _format_sfr_label(sfr: SFR, sfr_display_map: dict[str, dict[str, str]]) -> str:
    return sfr.sfr_id


def _describe_dependency_satisfaction(
    dependency_raw: str,
    sfr_by_code: dict[str, SFR],
    sfr_display_map: dict[str, dict[str, str]],
) -> str:
    groups = _extract_dependency_groups(dependency_raw)
    if not groups:
        return ""

    selected_labels: list[str] = []
    ambiguous_segments: list[str] = []

    for group in groups:
        matched_sfrs: list[SFR] = []
        for dependency_id in group:
            matched = sfr_by_code.get(dependency_id.upper())
            if matched and all(existing.id != matched.id for existing in matched_sfrs):
                matched_sfrs.append(matched)

        if not matched_sfrs:
            continue

        labels = [_format_sfr_label(item, sfr_display_map) for item in matched_sfrs]
        if len(group) == 1 or len(labels) == 1:
            selected_labels.append(labels[0])
        else:
            ambiguous_segments.append(" OR ".join(labels))

    parts: list[str] = []
    if selected_labels:
        parts.append(" AND ".join(selected_labels))
    if ambiguous_segments:
        parts.append("; OR-group satisfied by " + "; ".join(ambiguous_segments))
    return "".join(parts)


def _render_template(template: str, data: dict) -> str:
    """Replace {placeholder} tokens in template with actual data."""
    toe = data["toe"]
    assets = data["assets"]
    threats = data["threats"]
    assumptions = data["assumptions"]
    osps = data["osps"]
    objectives = data["objectives"]
    sfrs = data["sfrs"]
    tests = data["tests"]
    threat_objectives = data["threat_objectives"]
    assumption_objectives = data["assumption_objectives"]
    osp_objectives = data["osp_objectives"]
    objective_sfrs = data["objective_sfrs"]
    libraries_by_id = data["libraries_by_id"]
    libraries_by_component = data["libraries_by_component"]

    # Build lookup maps
    obj_map = {o.id: o for o in objectives}
    sfr_map = {s.id: s for s in sfrs}
    threat_map = {t.id: t for t in threats}
    assumption_map = {a.id: a for a in assumptions}
    osp_map = {o.id: o for o in osps}
    sfr_display_map = {
        sfr.id: _resolve_export_sfr_fields(sfr, libraries_by_id, libraries_by_component)
        for sfr in sfrs
    }

    # Simple field replacements
    replacements = {
        "TOE Name": toe.name or "—",
        "TOE Version": toe.version or "1.0",
        "TOE Type": toe.toe_type or "—",
        "TOE Brief Intro": _v(toe.brief_intro),
        "TOE Type Desc": _v(toe.toe_type_desc),
        "TOE Usage": _v(toe.toe_usage),
        "Major Security Features": _v(toe.major_security_features),
        "Required Non-TOE": _v(toe.required_non_toe_hw_sw_fw),
        "Physical Scope": _v(toe.physical_scope),
        "Logical Scope": _v(toe.logical_scope),
        "HW Interfaces": _v(toe.hw_interfaces),
        "SW Interfaces": _v(toe.sw_interfaces),
    }

    # Table replacements
    replacements["Assets Table"] = _md_table(
        ["Asset Name", "Asset Description"],
        [[a.name, _v(a.description)] for a in assets],
    )
    replacements["Threats Table"] = _md_table(
        ["Threat Code", "Threat Definition"],
        [[t.code, _threat_definition(t)] for t in threats],
    )
    replacements["Assumptions Table"] = _md_table(
        ["Assumption Code", "Assumption Definition"],
        [[a.code, _v(a.description)] for a in assumptions],
    )
    replacements["OSP Table"] = _md_table(
        ["OSP Code", "OSP Definition"],
        [[o.code, _v(o.description)] for o in osps],
    )
    obj_o, obj_oe = _objective_groups(objectives)
    replacements["Objectives Table"] = _md_table(
        ["Objective", "Objective Definition"],
        [[o.code, _v(o.description)] for o in obj_o],
    )
    replacements["OE Table"] = _md_table(
        ["Objective", "Objective Definition"],
        [[o.code, _v(o.description)] for o in obj_oe],
    )
    replacements["SFR Table"] = _md_sfr_sections(sfrs, sfr_display_map)
    replacements["Test Cases Table"] = _md_table(
        ["Case #", "Title", "Type", "Objective", "Steps", "Expected Result"],
        [[tc.case_number or "—", tc.title, tc.test_type or "—", _v(tc.objective), _v(tc.test_steps), _v(tc.expected_result)] for tc in tests],
    )

    # Field-level replacements
    replacements.update({
        "Asset Name": _md_list([a.name or "—" for a in assets]),
        "Asset Type": _md_list([a.asset_type or "—" for a in assets]),
        "Asset Importance": _md_list([str(a.importance) if a.importance is not None else "—" for a in assets]),
        "Asset Description": _md_list([_v(a.description) for a in assets]),
        "Threat Code": _md_list([t.code or "—" for t in threats]),
        "Threat Agent": _md_list([_v(t.threat_agent) for t in threats]),
        "Threat Definition": _md_list([_threat_definition(t) for t in threats]),
        "Threat Adverse Action": _md_list([_v(t.adverse_action) for t in threats]),
        "Threat Assets Affected": _md_list([_v(t.assets_affected) for t in threats]),
        "Threat Likelihood": _md_list([t.likelihood or "—" for t in threats]),
        "Threat Impact": _md_list([t.impact or "—" for t in threats]),
        "Threat Risk Level": _md_list([t.risk_level or "—" for t in threats]),
        "Assumption Code": _md_list([a.code or "—" for a in assumptions]),
        "Assumption Description": _md_list([_v(a.description) for a in assumptions]),
        "Assumption Definition": _md_list([_v(a.description) for a in assumptions]),
        "OSP Code": _md_list([o.code or "—" for o in osps]),
        "OSP Description": _md_list([_v(o.description) for o in osps]),
        "OSP Definition": _md_list([_v(o.description) for o in osps]),
        "Objective Code": _md_list([o.code or "—" for o in obj_o]),
        "Objective Description": _md_list([_v(o.description) for o in obj_o]),
        "Objective": _md_list([o.code or "—" for o in obj_o]),
        "Objective Definition": _md_list([_v(o.description) for o in obj_o]),
        "Objective Rationale": _md_list([_v(o.rationale) for o in obj_o]),
        "OE Code": _md_list([o.code or "—" for o in obj_oe]),
        "OE Description": _md_list([_v(o.description) for o in obj_oe]),
        "OE Definition": _md_list([_v(o.description) for o in obj_oe]),
        "OE Rationale": _md_list([_v(o.rationale) for o in obj_oe]),
        "SFR ID": _md_list([s.sfr_id or "—" for s in sfrs]),
        "SFR Name": _md_list([sfr_display_map[s.id]["name"] for s in sfrs]),
        "SFR Source": _md_list([s.source or "—" for s in sfrs]),
        "SFR Detail": _md_list([sfr_display_map[s.id]["detail"] for s in sfrs]),
        "SFR Customization": _md_list([_v(s.customization_note) for s in sfrs]),
        "Test Case Number": _md_list([tc.case_number or "—" for tc in tests]),
        "Test Case Title": _md_list([tc.title or "—" for tc in tests]),
        "Test Type": _md_list([tc.test_type or "—" for tc in tests]),
        "Test Objective": _md_list([_v(tc.objective) for tc in tests]),
        "Test Steps": _md_list([_v(tc.test_steps) for tc in tests]),
        "Test Expected Result": _md_list([_v(tc.expected_result) for tc in tests]),
    })

    # Mapping tables
    # Threat -> Objective
    t_o_rows = []
    for to in threat_objectives:
        t = threat_map.get(to.threat_id)
        o = obj_map.get(to.objective_id)
        if t and o:
            t_o_rows.append([t.code, o.code, _v(o.description)])
    replacements["Threat Objective Mapping"] = _md_table(
        ["Threat", "Objective", "Objective Description"], t_o_rows,
    )
    replacements["Threat Mapping Threat"] = _md_list([row[0] for row in t_o_rows])
    replacements["Threat Mapping Objective"] = _md_list([row[1] for row in t_o_rows])
    replacements["Threat Mapping Objective Description"] = _md_list([row[2] for row in t_o_rows])

    # Assumption -> Objective
    a_o_rows = []
    for ao in assumption_objectives:
        a = assumption_map.get(ao.assumption_id)
        o = obj_map.get(ao.objective_id)
        if a and o:
            a_o_rows.append([a.code, o.code, _v(o.description)])
    replacements["Assumption Objective Mapping"] = _md_table(
        ["Assumption", "Objective", "Objective Description"], a_o_rows,
    )
    replacements["Assumption Mapping Assumption"] = _md_list([row[0] for row in a_o_rows])
    replacements["Assumption Mapping Objective"] = _md_list([row[1] for row in a_o_rows])
    replacements["Assumption Mapping Objective Description"] = _md_list([row[2] for row in a_o_rows])

    # OSP -> Objective
    osp_o_rows = []
    for oo in osp_objectives:
        p = osp_map.get(oo.osp_id)
        o = obj_map.get(oo.objective_id)
        if p and o:
            osp_o_rows.append([p.code, o.code, _v(o.description)])
    replacements["OSP Objective Mapping"] = _md_table(
        ["OSP", "Objective", "Objective Description"], osp_o_rows,
    )
    replacements["OSP Mapping OSP"] = _md_list([row[0] for row in osp_o_rows])
    replacements["OSP Mapping Objective"] = _md_list([row[1] for row in osp_o_rows])
    replacements["OSP Mapping Objective Description"] = _md_list([row[2] for row in osp_o_rows])

    threat_objective_pairs = {(to.threat_id, to.objective_id) for to in threat_objectives}
    assumption_objective_pairs = {(ao.assumption_id, ao.objective_id) for ao in assumption_objectives}
    osp_objective_pairs = {(oo.osp_id, oo.objective_id) for oo in osp_objectives}

    tao_headers = ["Objective"] + [t.code for t in threats] + [a.code for a in assumptions] + [o.code for o in osps]
    tao_rows: list[list[str]] = []
    for objective in [*obj_o, *obj_oe]:
        row = [objective.code]
        row.extend("*" if (threat.id, objective.id) in threat_objective_pairs else "" for threat in threats)
        row.extend("*" if (assumption.id, objective.id) in assumption_objective_pairs else "" for assumption in assumptions)
        row.extend("*" if (osp.id, objective.id) in osp_objective_pairs else "" for osp in osps)
        tao_rows.append(row)
    replacements["TAO Objective Mapping"] = _md_table(tao_headers, tao_rows)

    replacements["Threat Objective rationale"] = _md_table(
        ["Threat Code", "Threat rationale"],
        [[t.code, _blank(t.ai_rationale)] for t in threats],
    )
    replacements["Assumption Objective rationale"] = _md_table(
        ["Assumption Code", "Assumption rationale"],
        [[a.code, ""] for a in assumptions],
    )
    replacements["OSP Objective rationale"] = _md_table(
        ["OSP Code", "OSP rationale"],
        [[o.code, ""] for o in osps],
    )

    # Objective -> SFR
    direct_objective_sfrs = [os_link for os_link in objective_sfrs if not _is_dependency_mapping(os_link.mapping_rationale)]
    direct_sfr_ids = {os_link.sfr_id for os_link in direct_objective_sfrs}
    direct_sfrs = [s for s in sfrs if s.id in direct_sfr_ids]
    objective_sfr_pairs = {(os_link.objective_id, os_link.sfr_id) for os_link in direct_objective_sfrs}
    o_s_rows = []
    for os_link in direct_objective_sfrs:
        o = obj_map.get(os_link.objective_id)
        s = sfr_map.get(os_link.sfr_id)
        if o and s:
            o_s_rows.append([o.code, s.sfr_id, sfr_display_map[s.id]["name"]])
    replacements["Objective SFR Mapping"] = _md_table(
        ["SFR"] + [o.code for o in obj_o],
        [[s.sfr_id] + ["*" if (o.id, s.id) in objective_sfr_pairs else "" for o in obj_o] for s in direct_sfrs],
    )
    replacements["Objective SFR Mapping Objective"] = _md_list([row[0] for row in o_s_rows])
    replacements["Objective SFR Mapping SFR"] = _md_list([row[1] for row in o_s_rows])
    replacements["Objective SFR Mapping Rationale"] = _md_list([f"{row[1]} {row[2]}" for row in o_s_rows])
    objective_to_direct_sfrs: dict[str, list[str]] = {}
    for os_link in direct_objective_sfrs:
        objective = obj_map.get(os_link.objective_id)
        sfr = sfr_map.get(os_link.sfr_id)
        if not objective or not sfr:
            continue
        objective_to_direct_sfrs.setdefault(objective.code, []).append(f"{sfr.sfr_id} {sfr_display_map[sfr.id]['name']}")
    replacements["Objective SFR rationale"] = _md_table(
        ["Security Objective", "SFR rationale"],
        [[objective.code, "; ".join(objective_to_direct_sfrs.get(objective.code, []))] for objective in obj_o],
    )

    sfr_id_lookup = {s.sfr_id.upper(): s for s in sfrs if s.sfr_id}
    sfr_dependency_rows: list[list[str]] = []
    for sfr in sfrs:
        dependency_raw = sfr_display_map[sfr.id]["dependency_raw"]
        if not dependency_raw:
            continue
        satisfaction = _describe_dependency_satisfaction(
            dependency_raw,
            sfr_id_lookup,
            sfr_display_map,
        )
        sfr_dependency_rows.append([
            sfr.sfr_id,
            dependency_raw,
            satisfaction,
        ])
    replacements["SFR Dependencies"] = _md_table(
        ["SFR", "Dependency", "Current TOE SFR"],
        sfr_dependency_rows,
    )

    # Apply replacements
    result = template
    for key, value in replacements.items():
        result = result.replace("{" + key + "}", value)
    return result


async def _get_template(db: AsyncSession) -> str:
    """Get ST template from local markdown file, migrating legacy DB content if needed."""
    file_content = _read_st_template_file()
    if file_content is not None:
        return file_content

    res = await db.exec(select(SystemSetting).where(SystemSetting.key == "st_template"))
    row = res.first()
    if row and row.value and row.value.strip():
        _write_st_template_file(row.value)
        return row.value

    _write_st_template_file(DEFAULT_ST_TEMPLATE)
    return DEFAULT_ST_TEMPLATE


# ── API endpoints ────────────────────────────────────────────────

@router.get("/toes/{toe_id}/st-preview")
async def get_st_preview(
    toe_id: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Render the ST template with current TOE data and return markdown."""
    await _get_toe(toe_id, current_user, db)
    template = await _get_template(db)
    data = await _load_all_data(toe_id, db)
    rendered = _render_template(template, data)
    return ok(data={"content": rendered})


@router.get("/st-template")
async def get_st_template(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Get the current ST template (markdown)."""
    template = await _get_template(db)
    return ok(data={"content": template})


@router.put("/st-template")
async def update_st_template(
    payload: dict,
    current_user: User = Depends(get_current_admin),
    db: AsyncSession = Depends(get_db),
):
    """Update ST template."""
    content = payload.get("content", "")
    _write_st_template_file(content)
    return ok(msg="Template saved")


@router.get("/st-template/default")
async def get_default_st_template(
    current_user: User = Depends(get_current_user),
):
    """Get the built-in default ST template."""
    return ok(data={"content": DEFAULT_ST_TEMPLATE})


class ExportSTBody(BaseModel):
    format: str = "md"  # md | docx
    content: Optional[str] = None  # if provided, export this content directly


@router.post("/toes/{toe_id}/export-st")
async def export_st(
    toe_id: str,
    body: ExportSTBody,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Export ST document in markdown or docx format."""
    await _get_toe(toe_id, current_user, db)

    # Get content: either user-provided or freshly rendered
    if body.content and body.content.strip():
        md_content = body.content
    else:
        template = await _get_template(db)
        data = await _load_all_data(toe_id, db)
        md_content = _render_template(template, data)

    toe_res = await db.exec(select(TOE).where(TOE.id == toe_id))
    toe = toe_res.first()
    toe_name = (toe.name if toe else "TOE").replace(" ", "_")
    now = datetime.now(timezone.utc).strftime("%Y-%m-%d")

    export_dir = os.path.join(settings.storage_path, "exports", toe_id)
    os.makedirs(export_dir, exist_ok=True)

    fmt = body.format.lower()
    if fmt == "docx":
        file_path = os.path.join(export_dir, f"ST_{toe_name}_{now}.docx")
        _export_docx(md_content, file_path, toe.name if toe else "TOE")
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        file_path = os.path.join(export_dir, f"ST_{toe_name}_{now}.md")
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(md_content)
        media_type = "text/markdown"

    from app.core.uploads import assert_within_storage
    assert_within_storage(file_path)
    filename = os.path.basename(file_path)
    return FileResponse(path=file_path, filename=filename, media_type=media_type)


def _export_docx(md_content: str, file_path: str, title: str):
    """Convert markdown to a simple docx file."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        # Fallback: save as plain text in docx
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(md_content)
        return

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    lines = md_content.split("\n")
    table_buffer: list[str] = []
    in_table = False

    def flush_table():
        nonlocal table_buffer, in_table
        if not table_buffer:
            return
        # Parse markdown table
        rows_raw = [r.strip().strip("|").split("|") for r in table_buffer if r.strip() and not re.match(r'^\|[\s\-:|]+\|$', r.strip())]
        if len(rows_raw) < 1:
            table_buffer = []
            in_table = False
            return
        col_count = len(rows_raw[0])
        tbl = doc.add_table(rows=len(rows_raw), cols=col_count, style='Table Grid')
        for i, row_cells in enumerate(rows_raw):
            for j, cell_text in enumerate(row_cells):
                if j < col_count:
                    tbl.rows[i].cells[j].text = cell_text.strip()
        doc.add_paragraph()
        table_buffer = []
        in_table = False

    for line in lines:
        stripped = line.strip()

        # Table detection
        if stripped.startswith("|"):
            in_table = True
            table_buffer.append(stripped)
            continue
        elif in_table:
            flush_table()

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("---"):
            doc.add_paragraph("─" * 50)
        elif stripped.startswith("*（") or stripped.startswith("*("):
            doc.add_paragraph(stripped, style='Normal')
        elif stripped == "":
            doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            # Handle bold text
            parts = re.split(r'(\*\*.*?\*\*)', stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    if in_table:
        flush_table()

    doc.save(file_path)


# ── Legacy download endpoint (kept for backward compat) ──────────

@router.get("/toes/{toe_id}/export-st/download/{task_id}")
async def download_st(
    toe_id: str,
    task_id: str,
    current_user: User = Depends(get_current_user),
    db=Depends(get_db),
):
    from app.models.ai_task import AITask
    await _get_toe(toe_id, current_user, db)
    res = await db.exec(
        select(AITask).where(
            AITask.id == task_id,
            AITask.user_id == current_user.id,
        )
    )
    task = res.first()
    if not task:
        raise HTTPException(404, "Task not found")
    if task.status != "done" or not task.download_url:
        raise HTTPException(400, "The file has not been generated yet")
    from app.core.uploads import assert_within_storage
    file_path = task.download_url
    assert_within_storage(file_path)
    if not os.path.exists(file_path):
        raise HTTPException(404, "The file does not exist and may have been cleaned up")
    filename = os.path.basename(file_path)
    ext = filename.rsplit(".", 1)[-1].lower()
    media_types = {
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "html": "text/html",
        "md": "text/markdown",
    }
    return FileResponse(
        path=file_path,
        filename=filename,
        media_type=media_types.get(ext, "application/octet-stream"),
    )
